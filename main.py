import os
import urllib.request
import webbrowser
import sys
import time
import threading
import subprocess
import tempfile
import zipfile
import stat
import platform
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Cm
from datetime import datetime
import json
from PIL import Image, ImageEnhance
import ssl
import hashlib

# --- MACOS CRASH FIX ---
if sys.platform == "darwin":
    os.environ["PATH"] += os.pathsep + "/usr/local/bin" + os.pathsep + "/opt/homebrew/bin"

# --- MACOS SSL CERT FIX ---
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context
# --------------------------------------------------------

import fitz 
from google import genai
from google.genai import types
import pypandoc

try:
    import pillow_heif
    pillow_heif.register_heif_opener()
except ImportError:
    print("[!] CẢNH BÁO: Chưa cài đặt pillow-heif. Không thể đọc file HEIC.")

CURRENT_VERSION = "2.2.1"
GITHUB_API_URL = "https://api.github.com/repos/tozn607/pdfscan2word/releases/latest"
RELEASES_URL = "https://github.com/tozn607/pdfscan2word/releases"

def get_build_date():
    if getattr(sys, 'frozen', False):
        try:
            mtime = os.path.getmtime(sys.executable)
            return datetime.fromtimestamp(mtime)
        except: pass
    return datetime.now()

CONFIG_DIR = os.path.join(os.path.expanduser("~"), ".pdfscan2word")
API_KEY_FILE = os.path.join(CONFIG_DIR, "api_key.txt")
CONFIG_JSON_FILE = os.path.join(CONFIG_DIR, "config.json")
CHECKPOINT_DIR = os.path.join(CONFIG_DIR, "checkpoints")
if not os.path.exists(CHECKPOINT_DIR): os.makedirs(CHECKPOINT_DIR)


safety_config = [
    types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HARASSMENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
]

PROMPT_VN = r"""
Bạn là một chuyên gia số hóa và phục hồi tài liệu chuyên nghiệp. Dưới đây là hình ảnh scan của một trang tài liệu/giáo trình. 
Nhiệm vụ của bạn là trích xuất và phục hồi, làm sạch văn bản theo các quy tắc NGHIÊM NGẶT sau đây:
1. CƠ CHẾ XUỐNG DÒNG: 
   - BẮT BUỘC chèn MỘT DÒNG TRẮNG (Enter 2 lần) giữa các đoạn văn, tiêu đề, mục danh sách để tránh bị dồn chữ.
   - Chỉ nối dòng nếu dòng dưới là câu đứt đoạn của dòng trên.
2. ĐẶC BIỆT XỬ LÝ MỤC LỤC (QUAN TRỌNG): 
   - Vẫn giữ nguyên một dòng trắng (Enter 2 lần) giữa các mục.
   - BẮT BUỘC khôi phục sự Thụt lề (Indentation) bằng cách chèn `&emsp;` vào đầu dòng (Mục lớn không thụt, Cấp 1 thụt 1 `&emsp;`, Cấp 2 thụt 2 `&emsp;`...).
   - TUYỆT ĐỐI KHÔNG GÕ dải dấu chấm (`......`) nối giữa tên mục và số trang. 
   - BẮT BUỘC phải thay thế toàn bộ dải dấu chấm đó bằng MỘT KÝ TỰ TAB DUY NHẤT với mã HTML là `&#9;`. 
   - Cú pháp chuẩn: `&emsp;1.1. Các phong cách học tập&#9;4`
3. NGĂN TỰ ĐỘNG TẠO BULLET POINT: 
   - NẾU bản gốc có dấu gạch ngang (-) ở đầu dòng, BẮT BUỘC phải dùng dấu gạch chéo ngược để thoát (escape): Viết là `\- ` thay vì `- `.
   - Các mục đánh số (1., 2.) hoặc chữ cái (a., b.) thì giữ nguyên, KHÔNG chèn thêm gạch ngang.
4. ĐỊNH DẠNG: **In đậm** và *In nghiêng* đúng bản gốc.
5. ĐIỀN CHỮ THIẾU: Dựa vào ngữ cảnh chung để điền bù chữ khuất mép giấy. Xóa bỏ hoàn toàn ký tự rác.
6. LOẠI BỎ SỐ TRANG: TUYỆT ĐỐI KHÔNG ghi lại số trang ở lề trên/dưới cùng của trang.
7. XỬ LÝ CHÚ THÍCH (FOOTNOTE):
   - Đặt mốc `[^1]`, `[^2]`... sát ngay sau từ/câu được chú thích.
   - Ghi nội dung chú thích ở tận cùng của văn bản theo cú pháp: `[^1]: Nội dung chú thích...`
8. THỤT ĐẦU DÒNG ĐOẠN VĂN: Nếu đoạn văn trong bản gốc có lùi vào ở dòng đầu tiên, chèn `&emsp;&emsp;` vào đầu đoạn.
9. BẢNG BIỂU (TABLES) - KHÔNG DÙNG HTML:
   - BẮT BUỘC sử dụng cú pháp bảng Markdown tiêu chuẩn (dùng các dấu gạch đứng `|`). TUYỆT ĐỐI KHÔNG DÙNG mã HTML.
   - XỬ LÝ Ô GỘP (MERGED CELLS): Điền nội dung vào hàng đầu tiên của nhóm ô gộp. Các hàng bên dưới thuộc cùng nhóm thì ĐỂ TRỐNG (ví dụ: `| | Toán | 105 |`).
10. BỎ QUA CHỮ BÓNG MỜ (BLEED-THROUGH): Do tài liệu in trên giấy mỏng, sẽ có các dòng chữ in hằn từ mặt sau (thường bị lật ngược, màu xám, mờ nhạt) nằm lẩn khuất dưới nền giấy. TUYỆT ĐỐI BỎ QUA những chữ này. Chỉ đọc và trích xuất văn bản chính (màu đen sẫm, rõ nét, đọc xuôi).
Chỉ trả về văn bản bằng Markdown, không giải thích gì thêm.
"""

PROMPT_EN = r"""
You are a professional document digitization and restoration expert. Here is a scanned image of a document/textbook page. 
Your task is to extract, restore, and clean the text following these STRICT rules:
1. LINE BREAK MECHANISM: 
   - YOU MUST insert ONE BLANK LINE (Enter twice) between paragraphs, headings, and lists to prevent clumping.
   - Only join lines if the bottom line is a continuation of a broken sentence.
2. TABLE OF CONTENTS (CRUCIAL): 
   - Keep one blank line between items.
   - YOU MUST restore Indentation by inserting `&emsp;` at the beginning of the line.
   - ABSOLUTELY DO NOT type a dot strip (`......`) connecting the section name and page number.
   - YOU MUST replace the entire dot strip with a SINGLE TAB character using the HTML code `&#9;`. 
   - Standard syntax: `&emsp;1.1. Learning styles&#9;4`
3. PREVENT AUTO BULLET POINTS: 
   - If the original has a dash (-) at the start of a line, YOU MUST escape it: Write `\- ` instead of `- `.
   - Numbered (1., 2.) or lettered (a., b.) items remain unchanged, DO NOT insert dashes.
4. FORMATTING: Preserve **Bold** and *Italics* exactly as the original.
5. FILL MISSING TEXT: Use context to fill in text cut off at the edges. Remove all garbage characters.
6. REMOVE PAGE NUMBERS: ABSOLUTELY DO NOT transcribe page numbers at the top/bottom margins.
7. FOOTNOTES:
   - Place markers `[^1]`, `[^2]`... immediately after the annotated word/sentence.
   - Write the footnote content at the very end of the text using the syntax: `[^1]: Footnote content...`
8. PARAGRAPH INDENTATION: If a paragraph is indented on the first line in the original, insert `&emsp;&emsp;` at the beginning.
9. TABLES - NO HTML:
   - YOU MUST use standard Markdown table syntax (using vertical pipes `|`). ABSOLUTELY NO HTML code.
   - MERGED CELLS: Fill the content in the first row of the merged group. Leave the rows below in the same group BLANK (e.g., `| | Math | 105 |`).
10. IGNORE BLEED-THROUGH TEXT: Due to thin scanned paper, there is faint, mirrored, or grayish text bleeding through from the back of the page. ABSOLUTELY IGNORE this background text. Only extract the primary, sharp, forward-reading text.
Only return the text in Markdown, provide no additional explanations.
"""

STRINGS = {
    "VN": {
        "title": "Chuyển Ảnh sang Word",
        "toolbar": "Thanh công cụ:",
        "lang_switch": "🌐 Ngôn ngữ",
        "merge_pdf": "🖼️ GỘP ẢNH THÀNH PDF",
        "mode_single": "Chế độ Đơn (1 File PDF)",
        "mode_batch": "Chế độ Hàng loạt (Thư mục)",
        "api_key": "Google API Key:",
        "load_api": "Tải từ file .txt",
        "api_placeholder": "Nhập API Key...",
        "select_pdf": "Chọn File PDF",
        "input_pdf_ph": "Đường dẫn đến 1 file PDF...",
        "input_dir_btn": "Thư mục Input",
        "input_dir_ph": "Thư mục chứa nhiều file PDF...",
        "output_dir": "Thư mục Output",
        "output_ph": "Trống = Tự động lưu cùng nơi với Input",
        "solve_opt": "🤖 AI Giải bài tập",
        "cover_opt": "🖼️ Lưu riêng trang bìa",
        "merge_opt": "📖 Gộp cặp trang (Tăng tốc & Tiết kiệm)",
        "merge_desc": "Bỏ chọn nếu: 1) Ảnh gốc đã chứa 2 trang; hoặc 2) Chữ trong trang đơn quá dày đặc.",
        "start": "▶ BẮT ĐẦU XỬ LÝ",
        "stop": "⏹ DỪNG LẠI",
        "timer": "Thời gian xử lý: {0:02d}:{1:02d}",
        "timer_init": "Thời gian: 00:00",
        "log_ready": "[*] Ứng dụng đã sẵn sàng. Hãy chọn file/thư mục để bắt đầu.",
        "log_stop_cmd": "\n[!] NHẬN LỆNH DỪNG... Đang hủy bỏ (Vui lòng đợi 1-2 giây).",
        "err_api": "[!] LỖI: Vui lòng nhập API Key.",
        "err_input": "[!] LỖI: Đường dẫn Input không tồn tại.",
        "err_output": "[!] LỖI: Thư mục Output không tồn tại.",
        "msg_auto_out": "[*] Không chọn Output, tự động lưu cùng thư mục Input: {0}",
        "log_no_pdf": "\n[!] Không tìm thấy file PDF nào để xử lý!",
        "log_start_batch": "\n[>>>] BẮT ĐẦU: Sẽ xử lý {0} file PDF.",
        "log_split": "\n[FILE {0}/{1}] Đang tách: {2}...",
        "log_pandoc": "  [*] Đang tải trình biên dịch Word (chỉ chạy 1 lần)...",
        "log_err_read": "  [X] LỖI ĐỌC PDF: {0}",
        "log_save_cover": "  [+] Đã lưu ảnh bìa: {0}",
        "log_err_cover": "  [!] Lỗi khi lưu ảnh bìa: {0}",
        "log_merge_pages": "  [*] Đang gộp cặp trang (trái-phải) để tăng tốc độ quét...",
        "log_read_block": "    [>] Đang đọc khối ảnh {0}/{1}...",
        "log_reject": "      [!] LỖI BẢN QUYỀN: Google từ chối đọc trang {0}.",
        "log_err_attempt": "      [!] Lỗi (Thử {0}/{1}): {2}",
        "log_skip": "      [X] Bỏ qua trang {0}.",
        "log_format_err": "  [!] Lỗi khi định dạng file Word: {0}",
        "log_save_draft": "  [BẢN NHÁP] Đã lưu kết quả tại: {0}",
        "log_save_ok": "  [***] Đã lưu kết quả tại: {0}",
        "log_rescue": "  [+] Đã cứu dữ liệu dưới dạng Markdown tại: {0}",
        "log_cancel": "\n[⏹] TIẾN TRÌNH ĐÃ BỊ HỦY BỞI NGƯỜI DÙNG.",
        "log_done": "\n[✓] ĐÃ HOÀN TẤT!",
        "log_resume": "[*] Đang tiếp tục từ trang {0}/{1}...",
        "log_bypass_try": "      [!] Google từ chối (Lỗi bản quyền). Đang thử Bypass lách bộ lọc (Lần {0})...",
        "log_net_error": "      [!] Lỗi kết nối (Thử lại sau 5 giây): {0}",
        "status_retrying": "Lỗi mạng - Đang thử lại...",
        "btn_unfinished": "📋 CHƯA HOÀN THÀNH",
        "btn_delete": "Xóa bản lưu",
        "btn_up": "Di chuyển lên",
        "btn_down": "Di chuyển xuống",
        "history_title": "Danh sách tiến trình chưa xong",
        "resume_title": "Tiếp tục tiến trình",
        "resume_query": "Tìm thấy bản nháp của file này từ trước. Bạn có muốn tiếp tục từ trang {0} không?",
        "dev_footer": "Developed by @tozn607 | Version v{0} Build {1} | © {2}",
        "merge_title": "Tiện ích: Gộp ảnh thành PDF",
        "add_img": "Thêm ảnh",
        "up": "▲ Lên",
        "down": "▼ Xuống",
        "clear": "Xóa hết",
        "compress": "🗜️ Nén giảm dung lượng",
        "enhance": "✨ Làm sáng & Rõ chữ",
        "export_pdf": "XUẤT RA PDF",
        "exporting": "Đang xử lý ảnh, vui lòng đợi...",
        "merge_success": "--- THÀNH CÔNG ---\nĐã lưu PDF: {0}",
        "merge_err": "[!] LỖI: {0}",
        "update_title": "Thông báo Cập nhật",
        "update_msg": "Có phiên bản mới: v{0}\nPhiên bản hiện tại: v{1}\n\nBạn có muốn tải bản cập nhật về không?",
        "btn_yes": "Tải ngay",
        "btn_no": "Bỏ qua",
        "prompt_solve_ext": "\n\n6. YÊU CẦU ĐẶC BIỆC: Tài liệu này chứa các bài tập/câu hỏi. Bạn BẮT BUỘC phải đọc và TRẢ LỜI/GIẢI CHI TIẾT các bài tập đó. Hãy tạo một phần 'ĐÁP ÁN' riêng biệt và rõ ràng ở cuối tài liệu và viết đáp án ở đó.",
        "speed_label": "🚀 Tốc độ xử lý:",
        "speed_eco": "Tiết kiệm (Eco)",
        "speed_balanced": "Cân bằng (Balanced)",
        "speed_turbo": "Tối đa (Turbo)",
        "desc_eco": "An toàn nhất cho key miễn phí. Rất ít khi bị lỗi giới hạn (Rate limit).",
        "desc_balanced": "Xử lý nhanh hơn. Có nguy cơ nhỏ bị Google giới hạn tạm thời.",
        "desc_turbo": "Tốc độ cực nhanh. Khuyến nghị dùng cho Key trả phí (Google Cloud).",
        "btn_update": "🔄 CẬP NHẬT",
        "build_info": "Thông tin bản dựng",
        "update_menu_title": "Cập nhật ứng dụng",
        "repo_btn": "🌐 Truy cập Repository",
        "check_update_btn": "🔍 Kiểm tra cập nhật",
        "current_ver_label": "Phiên bản hiện tại: {0}",
        "latest_ver_label": "Phiên bản mới nhất: {0}",
        "build_date_label": "Ngày dựng: {0}",
        "up_to_date_msg": "Bạn đang sử dụng phiên bản mới nhất.",
        "update_searching": "Đang kiểm tra..."
    },
    "EN": {
        "title": "Image to Word Converter",
        "toolbar": "Toolbar:",
        "lang_switch": "🌐 Language",
        "merge_pdf": "🖼️ MERGE IMAGES TO PDF",
        "mode_single": "Single Mode (1 PDF file)",
        "mode_batch": "Batch Mode (Folder)",
        "api_key": "Google API Key:",
        "load_api": "Load from .txt",
        "api_placeholder": "Enter API Key...",
        "select_pdf": "Select PDF",
        "input_pdf_ph": "Path to a single PDF file...",
        "input_dir_btn": "Input Folder",
        "input_dir_ph": "Folder containing multiple PDFs...",
        "output_dir": "Output Folder",
        "output_ph": "Empty = Auto save next to Input",
        "solve_opt": "🤖 AI Solves Exercises",
        "cover_opt": "🖼️ Save cover separately",
        "merge_opt": "📖 Merge page pairs (Faster & Cheaper)",
        "merge_desc": "Uncheck if: 1) Photos already contain 2 pages; or 2) Text in a single page is too dense.",
        "start": "▶ START PROCESSING",
        "stop": "⏹ STOP",
        "timer": "Processing time: {0:02d}:{1:02d}",
        "timer_init": "Time: 00:00",
        "log_ready": "[*] Application ready. Please select a file/folder to start.",
        "log_stop_cmd": "\n[!] STOP COMMAND RECEIVED... Canceling tasks (Please wait 1-2s).",
        "err_api": "[!] ERROR: Please enter API Key.",
        "err_input": "[!] ERROR: Input path does not exist.",
        "err_output": "[!] ERROR: Output folder does not exist.",
        "msg_auto_out": "[*] Output not selected, auto saving to Input folder: {0}",
        "log_no_pdf": "\n[!] No PDF files found to process!",
        "log_start_batch": "\n[>>>] START: Will process {0} PDF files.",
        "log_split": "\n[FILE {0}/{1}] Extracting: {2}...",
        "log_pandoc": "  [*] Downloading Word compiler (runs only once)...",
        "log_err_read": "  [X] PDF READ ERROR: {0}",
        "log_save_cover": "  [+] Saved cover image: {0}",
        "log_err_cover": "  [!] Error saving cover: {0}",
        "log_merge_pages": "  [*] Merging page pairs (left-right) to speed up scanning...",
        "log_read_block": "    [>] Reading image block {0}/{1}...",
        "log_reject": "      [!] COPYRIGHT ERROR: Google refused to read page {0}.",
        "log_err_attempt": "      [!] Error (Attempt {0}/{1}): {2}",
        "log_skip": "      [X] Skipping page {0}.",
        "log_format_err": "  [!] Error formatting Word file: {0}",
        "log_save_draft": "  [DRAFT] Result saved at: {0}",
        "log_save_ok": "  [***] Result saved at: {0}",
        "log_rescue": "  [+] Rescued data as Markdown at: {0}",
        "log_cancel": "\n[⏹] PROCESS CANCELLED BY USER.",
        "log_done": "\n[✓] COMPLETED!",
        "log_resume": "[*] Resuming from page {0}/{1}...",
        "log_bypass_try": "      [!] Google refused (Copyright error). Attempting Bypass filter (Try {0})...",
        "log_net_error": "      [!] Connection Error (Retrying in 5s): {0}",
        "status_retrying": "Net Error - Retrying...",
        "btn_unfinished": "📋 UNFINISHED",
        "btn_delete": "Delete Draft",
        "btn_up": "Move Up",
        "btn_down": "Move Down",
        "history_title": "Unfinished Tasks List",
        "resume_title": "Resume Process",
        "resume_query": "A previous draft for this file was found. Do you want to resume from page {0}?",
        "dev_footer": "Developed by @tozn607 | Version v{0} Build {1} | © {2}",
        "merge_title": "Utility: Merge Images to PDF",
        "add_img": "Add Images",
        "up": "▲ Up",
        "down": "▼ Down",
        "clear": "Clear All",
        "compress": "🗜️ Compress size",
        "enhance": "✨ Enhance & Brighten",
        "export_pdf": "EXPORT TO PDF",
        "exporting": "Processing images, please wait...",
        "merge_success": "--- SUCCESS ---\nSaved PDF: {0}",
        "merge_err": "[!] ERROR: {0}",
        "update_title": "Update Notification",
        "update_msg": "New version available: v{0}\nCurrent version: v{1}\n\nDo you want to download the update?",
        "btn_yes": "Download Now",
        "btn_no": "Skip",
        "prompt_solve_ext": "\n\n6. SPECIAL REQUIREMENT: This document contains exercises/questions. You MUST read and ANSWER/SOLVE them in detail. Create a clear, separate 'ANSWERS' section at the end of the document and write your answers there.",
        "speed_label": "🚀 Processing Speed:",
        "speed_eco": "Eco Mode",
        "speed_balanced": "Balanced",
        "speed_turbo": "Turbo Mode",
        "desc_eco": "Safest for free keys. Very low risk of reaching Google's limits.",
        "desc_balanced": "Faster processing. Moderate risk of temporary limits on free keys.",
        "desc_turbo": "Extremely fast. Recommended for Paid API keys (Cloud) to avoid errors.",
        "btn_update": "🔄 UPDATE",
        "build_info": "Build Information",
        "update_menu_title": "App Update",
        "repo_btn": "🌐 Open Repository",
        "check_update_btn": "🔍 Check for Updates",
        "current_ver_label": "Current Version: {0}",
        "latest_ver_label": "Latest Version: {0}",
        "build_date_label": "Build Date: {0}",
        "up_to_date_msg": "You are using the latest version.",
        "update_searching": "Checking..."
    }
}


from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QLabel, QLineEdit, QPushButton, 
                             QCheckBox, QComboBox, QTextEdit, QFileDialog, 
                             QMessageBox, QDialog, QButtonGroup, QListWidget, QFrame,
                             QSizePolicy, QScrollArea, QAbstractItemView, QProgressBar, QGroupBox)
from PyQt6.QtCore import Qt, pyqtSignal, QThread, QTimer, QSize
from PyQt6.QtGui import QFont, QIcon, QColor, QPalette, QCursor

QSS = """
    QWidget {
        font-size: 14px;
    }
    QFrame#Card {
        background-color: rgba(128, 128, 128, 0.08);
        border-radius: 12px;
        border: 1px solid rgba(128, 128, 128, 0.2);
    }
    QPushButton {
        border-radius: 8px;
        padding: 8px 16px;
        font-weight: bold;
        background-color: rgba(128, 128, 128, 0.15);
        border: 1px solid rgba(128, 128, 128, 0.2);
    }
    QPushButton:hover {
        background-color: rgba(128, 128, 128, 0.25);
    }
    QPushButton:checked {
        background-color: #1f538d;
        color: white;
        border: none;
    }
    QPushButton#Primary {
        background-color: #1f538d;
        color: white;
        border: none;
    }
    QPushButton#Primary:hover { background-color: #14375e; }
    QPushButton#Primary:disabled { background-color: #566c85; color: #d0d0d0; }
    
    QPushButton#Danger {
        background-color: #a83232;
        color: white;
        border: none;
    }
    QPushButton#Danger:hover { background-color: #7a2121; }
    QPushButton#Danger:disabled { background-color: #633f3f; color: #d0d0d0; }
    
    QPushButton#Success {
        background-color: #2b7a4b;
        color: white;
        border: none;
    }
    QPushButton#Success:hover { background-color: #1e5c37; }
    
    QLineEdit, QComboBox, QTextEdit, QListWidget {
        background-color: rgba(128, 128, 128, 0.05);
        border: 1px solid rgba(128, 128, 128, 0.3);
        border-radius: 6px;
        padding: 6px;
    }
    QLineEdit:focus, QComboBox:focus, QTextEdit:focus, QListWidget:focus {
        border: 1px solid #1f538d;
    }
    QProgressBar#Error {
        color: white;
    }
    QProgressBar#Error::chunk {
        background-color: #a83232;
    }
"""


class LanguageSelectorPopup(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Language / Chọn ngôn ngữ")
        self.setFixedSize(450, 300)
        self.selected_lang = "EN"
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        
        lbl = QLabel("Welcome! Please choose your preferred language.\nChào mừng! Vui lòng chọn ngôn ngữ của bạn.")
        font = QFont(".AppleSystemUIFont", 13, QFont.Weight.Bold)
        lbl.setFont(font)
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(lbl)
        
        layout.addSpacing(20)
        
        self.btn_en = QPushButton("🇺🇸 English")
        self.btn_en.setCheckable(True)
        self.btn_en.setChecked(True)
        self.btn_en.setFont(QFont(".AppleSystemUIFont", 12))
        
        self.btn_vn = QPushButton("🇻🇳 Tiếng Việt")
        self.btn_vn.setCheckable(True)
        self.btn_vn.setFont(QFont(".AppleSystemUIFont", 12))
        
        self.group = QButtonGroup(self)
        self.group.addButton(self.btn_en)
        self.group.addButton(self.btn_vn)
        
        layout.addWidget(self.btn_en)
        layout.addWidget(self.btn_vn)
        
        layout.addStretch()
        
        self.btn_save = QPushButton("Continue / Tiếp tục")
        self.btn_save.setObjectName("Success")
        self.btn_save.setMinimumHeight(45)
        self.btn_save.setFont(QFont(".AppleSystemUIFont", 13, QFont.Weight.Bold))
        self.btn_save.clicked.connect(self.on_save)
        layout.addWidget(self.btn_save)
        
    def on_save(self):
        if self.btn_vn.isChecked():
            self.selected_lang = "VN"
        else:
            self.selected_lang = "EN"
        self.accept()

class CheckpointHistoryDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.app = parent
        self.setWindowTitle(self.app.t("history_title"))
        self.setMinimumSize(750, 450)
        self.setStyleSheet(QSS)
        
        main_layout = QVBoxLayout(self)
        
        # Nội dung chính gồm List và Sidebar
        content_layout = QHBoxLayout()
        
        self.listbox = QListWidget()
        self.listbox.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.listbox.itemDoubleClicked.connect(self.on_select)
        content_layout.addWidget(self.listbox, 4)
        
        # Sidebar điều khiển
        sidebar = QVBoxLayout()
        self.btn_up = QPushButton(self.app.t("btn_up"))
        self.btn_up.clicked.connect(self.move_up)
        
        self.btn_down = QPushButton(self.app.t("btn_down"))
        self.btn_down.clicked.connect(self.move_down)
        
        self.btn_delete = QPushButton(self.app.t("btn_delete"))
        self.btn_delete.setStyleSheet("background-color: #a83232; color: white;")
        self.btn_delete.clicked.connect(self.delete_selected)
        
        sidebar.addWidget(self.btn_up)
        sidebar.addWidget(self.btn_down)
        sidebar.addStretch()
        sidebar.addWidget(self.btn_delete)
        
        content_layout.addLayout(sidebar, 1)
        main_layout.addLayout(content_layout)
        
        # Nút hành động chính
        actions = QHBoxLayout()
        self.btn_select = QPushButton("Resume / Tiếp tục")
        self.btn_select.setObjectName("Success")
        self.btn_select.setMinimumHeight(45)
        self.btn_select.clicked.connect(self.on_select)
        actions.addWidget(self.btn_select)
        main_layout.addLayout(actions)
        
        self.checkpoints = []
        self.load_checkpoints()
        
    def load_checkpoints(self):
        self.listbox.clear()
        self.checkpoints = []
        if not os.path.exists(CHECKPOINT_DIR): return
        
        files = [f for f in os.listdir(CHECKPOINT_DIR) if f.endswith(".json")]
        # Sắp xếp theo thời gian file modify mặc định
        files.sort(key=lambda x: os.path.getmtime(os.path.join(CHECKPOINT_DIR, x)), reverse=True)
        
        for f in files:
            try:
                p = os.path.join(CHECKPOINT_DIR, f)
                with open(p, "r", encoding="utf-8") as file:
                    data = json.load(file)
                    data["_internal_filename"] = f
                    filename = os.path.basename(data.get("pdf_path", "Unknown"))
                    page = data.get("current_page", 0)
                    total = data.get("total_pages", "?")
                    time_str = data.get("timestamp", "No Date")
                    
                    self.checkpoints.append(data)
                    display_text = f"📄 {filename}\n   ↳ {self.app.t('log_resume', page, total)} | {time_str}"
                    self.listbox.addItem(display_text)
            except: pass
                
    def move_up(self):
        row = self.listbox.currentRow()
        if row > 0:
            # Swap in memory
            self.checkpoints[row], self.checkpoints[row-1] = self.checkpoints[row-1], self.checkpoints[row]
            self.refresh_list()
            self.listbox.setCurrentRow(row - 1)

    def move_down(self):
        row = self.listbox.currentRow()
        if row < self.listbox.count() - 1 and row != -1:
            self.checkpoints[row], self.checkpoints[row+1] = self.checkpoints[row+1], self.checkpoints[row]
            self.refresh_list()
            self.listbox.setCurrentRow(row + 1)

    def delete_selected(self):
        row = self.listbox.currentRow()
        if row >= 0:
            data = self.checkpoints[row]
            f = data.get("_internal_filename")
            p = os.path.join(CHECKPOINT_DIR, f)
            try:
                if os.path.exists(p): os.remove(p)
                self.load_checkpoints()
            except: pass

    def refresh_list(self):
        self.listbox.clear()
        for data in self.checkpoints:
            filename = os.path.basename(data.get("pdf_path", "Unknown"))
            page = data.get("current_page", 0)
            total = data.get("total_pages", "?")
            time_str = data.get("timestamp", "No Date")
            display_text = f"📄 {filename}\n   ↳ {self.app.t('log_resume', page, total)} | {time_str}"
            self.listbox.addItem(display_text)

    def on_select(self):
        row = self.listbox.currentRow()
        if row >= 0:
            self.selected_data = self.checkpoints[row]
            self.accept()
        else:
            self.reject()

class WorkerThread(QThread):
    log_signal = pyqtSignal(str)
    progress_signal = pyqtSignal(int, int) # current, total
    status_signal = pyqtSignal(str, bool) # msg, is_error
    finished_signal = pyqtSignal()

    SPEED_CONFIGS = {
        0: {"threads": 1, "delay": 2.0},  # Eco
        1: {"threads": 3, "delay": 0.5},  # Balanced
        2: {"threads": 6, "delay": 0.0}   # Turbo
    }

    def __init__(self, app_instance, input_path, output_dir, mode, resume_at=0, initial_content="", speed_idx=1, client=None):
        super().__init__()
        self.app = app_instance
        self.client = client
        self.input_path = input_path
        self.output_dir = output_dir
        self.mode = mode
        self.resume_at = resume_at
        self.initial_content = initial_content
        self.speed_idx = speed_idx
        self.config = self.SPEED_CONFIGS.get(speed_idx, self.SPEED_CONFIGS[1])
        self.checkpoint_lock = threading.Lock()
        
    def run(self):
        import time 
        try:
            self.process_documents(self.input_path, self.output_dir, self.mode)
        except Exception as e:
            self.log_signal.emit(f"[!] ERROR: {str(e)}")
        finally:
            self.finished_signal.emit()

    def write_log(self, msg):
        self.log_signal.emit(msg)

    def get_checkpoint_path(self, pdf_path):
        import hashlib
        # Dùng absolute path để tránh lỗi trùng hash khác nhau khi restart app
        abs_path = os.path.abspath(pdf_path)
        pdf_id = hashlib.md5(abs_path.encode('utf-8')).hexdigest()
        return os.path.join(CHECKPOINT_DIR, f"{pdf_id}.json")

    def save_checkpoint(self, pdf_path, page_idx, content, total_pages=0):
        try:
            import datetime
            cp_path = self.get_checkpoint_path(pdf_path)
            data = {
                "pdf_path": pdf_path,
                "current_page": page_idx,
                "total_pages": total_pages,
                "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "content": content,
                "options": {
                    "solve": self.app.solve_var,
                    "cover": self.app.cover_var,
                    "merge": self.app.merge_pages_var,
                    "lang": self.app.current_lang,
                    "mode": self.app.t("mode_single") if self.app.btn_mode_single.isChecked() else self.app.t("mode_batch")
                }
            }
            with open(cp_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except: pass

    def clear_checkpoint(self, pdf_path):
        try:
            cp_path = self.get_checkpoint_path(pdf_path)
            if os.path.exists(cp_path):
                os.remove(cp_path)
        except: pass

    def transform_image_bypass(self, image, attempt):
        try:
            if attempt == 0: return image
            if attempt == 1: return ImageEnhance.Contrast(image).enhance(1.1)
            if attempt == 2: return ImageEnhance.Brightness(image).enhance(0.9)
            if attempt == 3:
                w, h = image.size
                return image.crop((2, 2, w - 2, h - 2))
            if attempt == 4: return ImageEnhance.Sharpness(image).enhance(1.5)
            return ImageEnhance.Contrast(image).enhance(1.2)
        except: return image

    def process_documents(self, input_path, output_dir, mode):
        if mode == self.app.t("mode_single"):
            pdf_files = [input_path]
        else:
            pdf_files = [os.path.join(input_path, f) for f in os.listdir(input_path) if f.lower().endswith('.pdf')]
        
        if not pdf_files:
            self.write_log(self.app.t("log_no_pdf"))
            return

        total_files = len(pdf_files)
        self.write_log(self.app.t("log_start_batch", total_files))
        model_id = 'gemini-3.1-flash-lite-preview'

        active_prompt = PROMPT_EN if self.app.current_lang == "EN" else PROMPT_VN
        if self.app.solve_var:
            active_prompt += self.app.t("prompt_solve_ext")

        for file_idx, pdf_path in enumerate(pdf_files):
            if self.app.stop_event.is_set(): break

            pdf_filename = os.path.basename(pdf_path)
            base_name = os.path.splitext(pdf_filename)[0]
            suffix = "_hoanthien" if self.app.current_lang == "VN" else "_completed"
            output_docx_path = os.path.join(output_dir, f"{base_name}{suffix}.docx")

            self.write_log(self.app.t("log_split", file_idx + 1, total_files, pdf_filename))
            
            images = []
            try: 
                pandoc_exe = os.path.join(CONFIG_DIR, "pandoc" + (".exe" if sys.platform == "win32" else ""))
                if os.path.exists(pandoc_exe):
                    os.environ['PYPANDOC_PANDOC'] = pandoc_exe

                try:
                    pypandoc.get_pandoc_version()
                except OSError:
                    self.write_log(self.app.t("log_pandoc"))
                    pypandoc.download_pandoc(targetfolder=CONFIG_DIR, download_folder=CONFIG_DIR)
                    os.environ['PYPANDOC_PANDOC'] = pandoc_exe

                doc = fitz.open(pdf_path)
                for page in doc:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # --- BỘ LỌC TẨY BÓNG MỜ GIẤY MỎNG ---
                    # Tăng tương phản ép màu xám mờ biến mất vào nền trắng, chữ đen sắc nét hơn
                    img = ImageEnhance.Contrast(img).enhance(1.6)
                    # Tăng sáng nhẹ để tẩy các vùng tối màu ở mép sách
                    img = ImageEnhance.Brightness(img).enhance(1.1)
                    
                    images.append(img)
                    
            except Exception as e:
                self.write_log(self.app.t("log_err_read", e))
                continue 

            if not images:
                continue

            start_idx = 0
            if self.app.cover_var:
                cover_path = os.path.join(output_dir, f"{base_name}_cover.jpg")
                try:
                    images[0].save(cover_path, "JPEG")
                    self.write_log(self.app.t("log_save_cover", f"{base_name}_cover.jpg"))
                except Exception as e:
                    self.write_log(self.app.t("log_err_cover", e))
                start_idx = 1 

            processed_images = []
            if self.app.merge_pages_var:
                self.write_log(self.app.t("log_merge_pages"))
                for i in range(start_idx, len(images), 2):
                    img_left = images[i]
                    if i + 1 < len(images):
                        img_right = images[i+1]
                        total_width = img_left.width + img_right.width
                        max_height = max(img_left.height, img_right.height)
                        new_img = Image.new('RGB', (total_width, max_height))
                        new_img.paste(img_left, (0, 0))
                        new_img.paste(img_right, (img_left.width, 0))
                        processed_images.append(new_img)
                    else:
                        processed_images.append(img_left)
            else:
                processed_images = images[start_idx:]

            # --- RESUME LOGIC ---
            file_resume_at = 0
            file_initial_content = ""
            
            # Nếu là file đầu tiên và có resume_at từ constructor (Single Mode), ưu tiên dùng nó
            if file_idx == 0 and self.resume_at > 0:
                file_resume_at = self.resume_at
                file_initial_content = self.initial_content
                self.write_log(self.app.t("log_resume", file_resume_at, len(processed_images)))
            else:
                cp_path = self.get_checkpoint_path(pdf_path)
                if os.path.exists(cp_path):
                    try:
                        with open(cp_path, "r", encoding="utf-8") as f:
                            cp_data = json.load(f)
                            # Chỉ resume nếu các option giống hệt nhau
                            opts = cp_data.get("options", {})
                            if (opts.get("solve") == self.app.solve_var and 
                                opts.get("cover") == self.app.cover_var and 
                                opts.get("merge") == self.app.merge_pages_var):
                                
                                file_resume_at = cp_data.get("current_page", 0)
                                file_initial_content = cp_data.get("content", "")
                                if file_resume_at > 0:
                                    self.write_log(self.app.t("log_resume", file_resume_at, len(processed_images)))
                    except: pass

            full_markdown_content = file_initial_content
            max_retries = 10 
            total_blocks = len(processed_images)
            self.progress_signal.emit(file_resume_at, total_blocks)

            # --- PARALLEL PROCESSING ---
            from concurrent.futures import ThreadPoolExecutor, as_completed
            
            results_map = {}
            next_to_save = file_resume_at

            def process_page(idx):
                if self.app.stop_event.is_set(): return None
                img = processed_images[idx]
                self.write_log(self.app.t("log_read_block", idx + 1, total_blocks))
                
                for attempt in range(max_retries):
                    if self.app.stop_event.is_set(): return None
                    try:
                        # Thêm delay nhỏ trước mỗi request nếu ở chế độ Eco/Balanced để tránh burst rate
                        if self.config['delay'] > 0 and attempt == 0:
                            time.sleep(self.config['delay'] * (idx % self.config['threads']))

                        response = self.client.models.generate_content(
                            model=model_id,
                            contents=[active_prompt, img],
                            config=types.GenerateContentConfig(
                                safety_settings=safety_config
                            )
                        )
                        try:
                            text_result = response.text
                        except ValueError:
                            # Lỗi bản quyền / Safety filter
                            if attempt < 5:
                                transform_img = self.transform_image_bypass(img, attempt + 1)
                                response = self.client.models.generate_content(
                                    model=model_id,
                                    contents=[active_prompt, transform_img],
                                    config=types.GenerateContentConfig(
                                        safety_settings=safety_config
                                    )
                                )
                                text_result = response.text
                            else:
                                return (idx, f"\n\n\n\n> **[COPYRIGHT BLOCK: PAGE {idx + 1}]**\n\n")

                        text_result = text_result.replace("[^", f"[^p{idx}_")
                        return (idx, f"\n\n\n\n{text_result}\n\n")

                    except Exception as e:
                        msg = str(e)
                        is_conn_error = any(kw in msg.lower() for kw in ["connection", "network", "timeout", "http"])
                        is_rate_limit = any(kw in msg.lower() for kw in ["429", "quota", "too many requests"])

                        if is_rate_limit:
                            wait_time = (attempt + 1) * 10
                            self.write_log(f"      [!] Rate Limit (429). Waiting {wait_time}s...")
                            time.sleep(wait_time)
                        elif is_conn_error:
                            time.sleep(5)
                        else:
                            time.sleep(10)
                        
                        if attempt == max_retries - 1:
                            return (idx, f"\n\n> **[PAGE {idx + 1} ERROR: {str(e)}]**\n\n")
                return None

            with ThreadPoolExecutor(max_workers=self.config['threads']) as executor:
                futures = {executor.submit(process_page, i): i for i in range(file_resume_at, total_blocks)}
                for future in as_completed(futures):
                    if self.app.stop_event.is_set():
                        executor.shutdown(wait=False, cancel_futures=True)
                        break
                    
                    res = future.result()
                    if res:
                        p_idx, p_text = res
                        with self.checkpoint_lock:
                            results_map[p_idx] = p_text
                            while next_to_save in results_map:
                                full_markdown_content += results_map[next_to_save]
                                next_to_save += 1
                                self.save_checkpoint(pdf_path, next_to_save, full_markdown_content, total_blocks)
                                self.progress_signal.emit(next_to_save, total_blocks)

            if self.app.stop_event.is_set(): break

            if full_markdown_content.strip() and not self.app.stop_event.is_set():
                # Xóa checkpoint CHI KHI hoàn thành thực sự (không phải dừng)
                self.clear_checkpoint(pdf_path)

            if full_markdown_content.strip():
                try:
                    pypandoc.convert_text(full_markdown_content, 'docx', format='md', outputfile=output_docx_path)
                    try:
                        doc_docx = docx.Document(output_docx_path)
                        for paragraph in doc_docx.paragraphs:
                            if '\t' in paragraph.text:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                paragraph.paragraph_format.space_after = Pt(0)
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.tab_stops.clear_all()
                                paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        for table in doc_docx.tables:
                            tblPr = table._tbl.tblPr
                            tblBorders = OxmlElement('w:tblBorders')
                            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single') 
                                border.set(qn('w:sz'), '4')       
                                border.set(qn('w:space'), '0')    
                                border.set(qn('w:color'), '000000') 
                                tblBorders.append(border)
                            tblPr.append(tblBorders)
                            
                        doc_docx.save(output_docx_path)
                    except Exception as e_docx:
                        self.write_log(self.app.t("log_format_err", e_docx))

                    if self.app.stop_event.is_set():
                        self.write_log(self.app.t("log_save_draft", output_docx_path))
                    else:
                        self.write_log(self.app.t("log_save_ok", output_docx_path))
                except Exception as e:
                    md_path = output_docx_path.replace('.docx', '.md')
                    with open(md_path, 'w', encoding='utf-8') as f: 
                        f.write(full_markdown_content)
                    self.write_log(self.app.t("log_rescue", md_path))
            
            if self.app.stop_event.is_set():
                break

        if self.app.stop_event.is_set():
            self.write_log(self.app.t("log_cancel"))
        else:
            self.write_log(self.app.t("log_done"))

class MergeWindow(QDialog):
    def __init__(self, app_instance):
        super().__init__(app_instance)
        self.app = app_instance
        self.setWindowTitle(self.app.t("merge_title"))
        self.setMinimumSize(660, 520)
        self.setStyleSheet(QSS)
        
        self.selected_images_list = []
        layout = QVBoxLayout(self)
        
        controls = QHBoxLayout()
        self.btn_add = QPushButton(self.app.t("add_img"))
        self.btn_add.setObjectName("Success")
        self.btn_add.clicked.connect(self.add_images)
        
        self.btn_up = QPushButton(self.app.t("up"))
        self.btn_up.clicked.connect(self.move_up)
        
        self.btn_down = QPushButton(self.app.t("down"))
        self.btn_down.clicked.connect(self.move_down)
        
        self.btn_clear = QPushButton(self.app.t("clear"))
        self.btn_clear.setObjectName("Danger")
        self.btn_clear.clicked.connect(self.clear_images)
        
        controls.addWidget(self.btn_add)
        controls.addWidget(self.btn_up)
        controls.addWidget(self.btn_down)
        controls.addStretch()
        controls.addWidget(self.btn_clear)
        layout.addLayout(controls)
        
        opt_layout = QHBoxLayout()
        self.chk_compress = QCheckBox(self.app.t("compress"))
        self.chk_compress.setChecked(True)
        self.chk_enhance = QCheckBox(self.app.t("enhance"))
        opt_layout.addWidget(self.chk_compress)
        opt_layout.addWidget(self.chk_enhance)
        opt_layout.addStretch()
        layout.addLayout(opt_layout)
        
        self.listbox = QListWidget()
        self.listbox.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        layout.addWidget(self.listbox)
        
        self.btn_export = QPushButton(self.app.t("export_pdf"))
        self.btn_export.setObjectName("Success")
        self.btn_export.setMinimumHeight(50)
        self.btn_export.setFont(QFont(".AppleSystemUIFont", 14, QFont.Weight.Bold))
        self.btn_export.clicked.connect(self.export_to_pdf)
        layout.addWidget(self.btn_export)
        
    def add_images(self):
        paths, _ = QFileDialog.getOpenFileNames(self, "Select Images", "", "Image Files (*.png *.jpg *.jpeg *.bmp *.heic *.heif)")
        if paths:
            self.selected_images_list.extend(paths)
            self.update_listbox()
            
    def move_up(self):
        row = self.listbox.currentRow()
        if row > 0:
            self.selected_images_list[row-1], self.selected_images_list[row] = self.selected_images_list[row], self.selected_images_list[row-1]
            self.update_listbox()
            self.listbox.setCurrentRow(row-1)

    def move_down(self):
        row = self.listbox.currentRow()
        if 0 <= row < len(self.selected_images_list) - 1:
            self.selected_images_list[row+1], self.selected_images_list[row] = self.selected_images_list[row], self.selected_images_list[row+1]
            self.update_listbox()
            self.listbox.setCurrentRow(row+1)

    def clear_images(self):
        self.selected_images_list.clear()
        self.update_listbox()
        
    def update_listbox(self):
        self.listbox.clear()
        for p in self.selected_images_list:
            self.listbox.addItem(os.path.basename(p))

    def export_to_pdf(self):
        if not self.selected_images_list: return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "", "PDF (*.pdf)")
        if save_path:
            self.listbox.clear()
            self.listbox.addItem(self.app.t("exporting"))
            QApplication.processEvents()
            
            try:
                processed = []
                for p in self.selected_images_list:
                    img = Image.open(p).convert('RGB')
                    if self.chk_enhance.isChecked():
                        img = ImageEnhance.Contrast(img).enhance(1.5)
                        img = ImageEnhance.Sharpness(img).enhance(2.0)
                        img = ImageEnhance.Brightness(img).enhance(1.1)
                    if self.chk_compress.isChecked():
                        max_w = 1200
                        if img.width > max_w:
                            ratio = max_w / img.width
                            img = img.resize((max_w, int(img.height * ratio)), Image.Resampling.LANCZOS)
                    processed.append(img)
                processed[0].save(save_path, save_all=True, append_images=processed[1:])
                self.listbox.clear()
                self.listbox.addItem(self.app.t("merge_success", save_path))
            except Exception as e:
                self.listbox.clear()
                self.listbox.addItem(self.app.t("merge_err", str(e)))

class DownloadUpdateThread(QThread):
    progress_signal = pyqtSignal(int)
    log_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str, str)
    
    def __init__(self, download_url, version):
        super().__init__()
        self.download_url = download_url
        self.version = version
        
    def run(self):
        try:
            self.log_signal.emit(f"Đang tải bản cập nhật (v{self.version})...")
            
            temp_dir = tempfile.mkdtemp(prefix="pdfscan2word_update_")
            zip_path = os.path.join(temp_dir, "update.zip")
            
            req = urllib.request.Request(self.download_url, headers={'User-Agent': 'PDFScan2Word-App'})
            with urllib.request.urlopen(req) as response:
                total_size = int(response.info().get("Content-Length", -1))
                
                downloaded = 0
                with open(zip_path, 'wb') as f:
                    while True:
                        chunk = response.read(8192)
                        if not chunk:
                            break
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total_size > 0:
                            percent = int(downloaded * 100 / total_size)
                            self.progress_signal.emit(percent)
            
            self.log_signal.emit("Tải xong! Đang giải nén...")
            
            if sys.platform == "darwin":
                import subprocess
                subprocess.run(["ditto", "-xk", zip_path, temp_dir], check=True)
            else:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
            self.log_signal.emit("Sẵn sàng cài đặt...")
            self.finished_signal.emit(temp_dir, "")
        except Exception as e:
            self.finished_signal.emit("", str(e))

class UpdateProgressDialog(QDialog):
    def __init__(self, parent=None, download_url="", version=""):
        super().__init__(parent)
        self.setWindowTitle("Đang cập nhật...")
        self.setFixedSize(400, 150)
        
        layout = QVBoxLayout(self)
        
        self.lbl_status = QLabel("Đang kết nối...", self)
        layout.addWidget(self.lbl_status)
        
        self.progress = QProgressBar(self)
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        layout.addWidget(self.progress)
        
        self.thread = DownloadUpdateThread(download_url, version)
        self.thread.progress_signal.connect(self.progress.setValue)
        self.thread.log_signal.connect(self.lbl_status.setText)
        self.thread.finished_signal.connect(self.on_download_finished)
        self.thread.start()
        
    def on_download_finished(self, temp_dir, err):
        if err:
            QMessageBox.critical(self, "Lỗi cập nhật", f"Không thể tải/giải nén:\\n{err}")
            self.reject()
        else:
            self.parent().perform_update_swap(temp_dir)
            self.accept()

class UpdateMenuDialog(QDialog):
    check_finished_signal = pyqtSignal(bool)
    check_error_signal = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.app = parent
        self.setWindowTitle(self.app.t("update_menu_title"))
        self.setFixedSize(450, 320)
        self.setStyleSheet(QSS)

        self.check_finished_signal.connect(self.finish_check)
        self.check_error_signal.connect(self.handle_check_error)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(25, 25, 25, 25)
        layout.setSpacing(15)
        
        # Build Info Section
        info_group = QGroupBox(self.app.t("build_info"))
        info_group.setStyleSheet("QGroupBox { font-weight: bold; border: 1px solid #ccc; border-radius: 8px; margin-top: 10px; padding-top: 10px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }")
        info_layout = QVBoxLayout(info_group)
        
        self.lbl_current = QLabel(self.app.t("current_ver_label", CURRENT_VERSION))
        self.lbl_current.setFont(QFont(".AppleSystemUIFont", 11))
        
        b_date = self.app.build_date.strftime("%Y-%m-%d %H:%M:%S")
        self.lbl_build_date = QLabel(self.app.t("build_date_label", f"{b_date} ({self.app.build_str})"))
        self.lbl_build_date.setFont(QFont(".AppleSystemUIFont", 11))
        
        info_layout.addWidget(self.lbl_current)
        info_layout.addWidget(self.lbl_build_date)
        layout.addWidget(info_group)
        
        # Latest Version Info
        self.lbl_latest = QLabel(self.app.t("latest_ver_label", self.app.latest_version if self.app.latest_version else "---"))
        self.lbl_latest.setFont(QFont(".AppleSystemUIFont", 11, QFont.Weight.Bold))
        layout.addWidget(self.lbl_latest)
        
        # Status Label
        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: #3a7ebf;")
        self.lbl_status.setWordWrap(True)
        layout.addWidget(self.lbl_status)
        
        # Actions
        actions_layout = QHBoxLayout()
        self.btn_check = QPushButton(self.app.t("check_update_btn"))
        self.btn_check.setMinimumHeight(35)
        self.btn_check.clicked.connect(self.check_now)
        
        self.btn_repo = QPushButton(self.app.t("repo_btn"))
        self.btn_repo.setMinimumHeight(35)
        self.btn_repo.clicked.connect(self.open_repo)
        
        actions_layout.addWidget(self.btn_check)
        actions_layout.addWidget(self.btn_repo)
        layout.addLayout(actions_layout)
        
        self.btn_install = QPushButton(self.app.t("btn_yes")) 
        self.btn_install.setObjectName("Success")
        self.btn_install.setMinimumHeight(45)
        self.btn_install.setVisible(False)
        self.btn_install.clicked.connect(self.install_update)
        layout.addWidget(self.btn_install)
        
        if self.app.latest_version and self.app.latest_version > CURRENT_VERSION:
            self.show_update_available()

    def check_now(self):
        self.lbl_status.setText(self.app.t("update_searching"))
        self.btn_check.setEnabled(False)
        threading.Thread(target=self.run_check, daemon=True).start()
        
    def run_check(self):
        try:
            req = urllib.request.Request(GITHUB_API_URL, headers={'User-Agent': 'PDFScan2Word-App'})
            with urllib.request.urlopen(req, timeout=8) as r:
                data = json.loads(r.read().decode('utf-8'))
                latest = data.get('tag_name', '').lstrip('v')
            
            if not latest:
                raise Exception("Could not find tag_name in API response")

            # Fetch download URL if available
            download_url = ""
            is_new = False
            # Simple version comparison (works for strictly increasing numeric strings like 2.2.0)
            if latest > CURRENT_VERSION:
                is_new = True
                if sys.platform == "win32":
                    suffix = "Windows.zip"
                elif sys.platform == "darwin":
                    arch = platform.machine()
                    suffix = "macOS-arm64.zip" if arch == "arm64" else "macOS-x86_64.zip"
                else:
                    suffix = None
                
                if suffix:
                    for asset in data.get('assets', []):
                        if asset.get('name', '').endswith(suffix):
                            download_url = asset.get('browser_download_url', '')
                            break
            
            self.app.latest_version = latest
            self.app.download_url = download_url
            
            self.check_finished_signal.emit(is_new)

        except Exception as e:
            self.check_error_signal.emit(str(e))

    def handle_check_error(self, msg):
        self.lbl_status.setText(f"Error: {msg}")
        self.btn_check.setEnabled(True)

    def finish_check(self, available):
        self.btn_check.setEnabled(True)
        self.lbl_latest.setText(self.app.t("latest_ver_label", self.app.latest_version))
        if available:
            self.show_update_available()
        else:
            self.lbl_status.setText(self.app.t("up_to_date_msg"))
            self.btn_install.setVisible(False)

    def show_update_available(self):
        self.lbl_status.setText(f"New version v{self.app.latest_version} is available!")
        self.btn_install.setVisible(True)

    def open_repo(self):
        webbrowser.open("https://github.com/tozn607/pdfscan2word")

    def install_update(self):
        if self.app.download_url:
            self.accept()
            self.app.prompt_update(self.app.latest_version, self.app.download_url)
        else:
            webbrowser.open(RELEASES_URL)
            self.accept()

class PDFOCRApp(QMainWindow):
    update_available_signal = pyqtSignal(str, str)
    
    def __init__(self):
        super().__init__()
        self.current_lang = "VN" # Default
        self.solve_var = False
        self.cover_var = False
        self.merge_pages_var = True
        self.speed_idx = 1 # Balanced by default
        self.load_config()
        self.stop_event = threading.Event()
        self.start_time = None
        self.timer_running = False

        self.setWindowTitle(self.t("title") + " v" + CURRENT_VERSION)
        self.setMinimumSize(700, 750)
        self.resize(840, 800)
        self.setStyleSheet(QSS)
        
        self.latest_version = None
        self.download_url = None
        self.build_date = get_build_date()
        self.build_str = self.build_date.strftime("%Y%m%d")
        self.year_str = self.build_date.strftime("%Y")

        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_timer_ui)
        
        self.worker = None

        self.build_ui()
        self.update_ui_texts()
        
        if not os.path.exists(CONFIG_JSON_FILE):
            QTimer.singleShot(200, self.show_language_popup)

        # Update connection
        self.update_available_signal.connect(self.prompt_update)
        # Update check threaded
        threading.Thread(target=self.check_for_updates, daemon=True).start()
        
        # Auto check for unfinished tasks on startup
        QTimer.singleShot(1000, self.auto_check_checkpoints)

    def t(self, key, *args):
        text = STRINGS.get(self.current_lang, STRINGS["VN"]).get(key, key)
        if args: return text.format(*args)
        return text

    def load_config(self):
        if os.path.exists(CONFIG_JSON_FILE):
            try:
                with open(CONFIG_JSON_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.current_lang = data.get("lang", "VN")
                    self.solve_var = data.get("solve", False)
                    self.cover_var = data.get("cover", False)
                    self.merge_pages_var = data.get("merge", True)
                    self.speed_idx = data.get("speed", 1)
            except: pass

    def save_config(self):
        if not os.path.exists(CONFIG_DIR): os.makedirs(CONFIG_DIR)
        try:
            with open(CONFIG_JSON_FILE, "w", encoding="utf-8") as f:
                json.dump({
                    "lang": self.current_lang,
                    "solve": self.solve_var,
                    "cover": self.cover_var,
                    "merge": self.merge_pages_var,
                    "speed": self.speed_idx
                }, f)
        except: pass

    def show_language_popup(self):
        dlg = LanguageSelectorPopup(self)
        if dlg.exec():
            self.current_lang = dlg.selected_lang
            self.save_config()
            idx = 0 if self.current_lang == "EN" else 1
            self.lang_menu.setCurrentIndex(idx)
            self.update_ui_texts()

    def build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # Menu frame
        menu_frame = QFrame()
        menu_frame.setStyleSheet("background-color: rgba(128, 128, 128, 0.12); border-bottom: 1px solid rgba(128, 128, 128, 0.2);")
        menu_layout = QHBoxLayout(menu_frame)
        menu_layout.setContentsMargins(20, 10, 20, 10)
        menu_layout.setSpacing(12)
        self.lbl_toolbar = QLabel("")
        self.lbl_toolbar.setFont(QFont(".AppleSystemUIFont", 12, QFont.Weight.Bold))
        self.btn_merge = QPushButton("")
        self.btn_merge.setMinimumHeight(40)
        self.btn_merge.clicked.connect(self.open_merge_popup)
        
        self.btn_unfinished = QPushButton("")
        self.btn_unfinished.setMinimumHeight(40)
        self.btn_unfinished.clicked.connect(self.open_unfinished_manager)
        
        self.btn_update = QPushButton("")
        self.btn_update.setMinimumHeight(40)
        self.btn_update.clicked.connect(self.open_update_menu)
        
        menu_layout.addWidget(self.lbl_toolbar)
        menu_layout.addWidget(self.btn_merge)
        menu_layout.addWidget(self.btn_unfinished)
        menu_layout.addWidget(self.btn_update)
        menu_layout.addStretch()
        
        self.lbl_lang = QLabel("🌐")
        self.lang_menu = QComboBox()
        self.lang_menu.addItems(["EN (English)", "VN (Tiếng Việt)"])
        idx = 0 if self.current_lang == "EN" else 1
        self.lang_menu.setCurrentIndex(idx)
        self.lang_menu.currentIndexChanged.connect(self.change_language)
        menu_layout.addWidget(self.lbl_lang)
        menu_layout.addWidget(self.lang_menu)
        main_layout.addWidget(menu_frame)

        # Content 
        content_layout = QVBoxLayout()
        content_layout.setContentsMargins(25, 20, 25, 20)
        content_layout.setSpacing(12)
        main_layout.addLayout(content_layout)

        # Segmented Button (Mode)
        mode_layout = QHBoxLayout()
        self.btn_mode_single = QPushButton("")
        self.btn_mode_single.setCheckable(True)
        self.btn_mode_single.setChecked(True)
        self.btn_mode_single.setMinimumHeight(45)
        self.btn_mode_batch = QPushButton("")
        self.btn_mode_batch.setCheckable(True)
        self.btn_mode_batch.setMinimumHeight(45)
        
        self.mode_group = QButtonGroup(self)
        self.mode_group.addButton(self.btn_mode_single)
        self.mode_group.addButton(self.btn_mode_batch)
        self.mode_group.buttonClicked.connect(self.change_mode)
        
        mode_layout.addWidget(self.btn_mode_single)
        mode_layout.addWidget(self.btn_mode_batch)
        content_layout.addLayout(mode_layout)
        # API KEY CARD
        api_card = QFrame(); api_card.setObjectName("Card")
        api_layout = QHBoxLayout(api_card)
        self.lbl_api = QLabel("")
        self.lbl_api.setFont(QFont(".AppleSystemUIFont", 11, QFont.Weight.Bold))
        self.entry_api = QLineEdit()
        self.entry_api.setEchoMode(QLineEdit.EchoMode.Password)
        self.entry_api.setMinimumHeight(40)
        self.btn_load_api = QPushButton("")
        self.btn_load_api.setMinimumHeight(40)
        self.btn_load_api.clicked.connect(self.load_api_from_file)
        
        api_layout.addWidget(self.lbl_api)
        api_layout.addWidget(self.entry_api)
        api_layout.addWidget(self.btn_load_api)
        content_layout.addWidget(api_card)
        # I/O CARD
        io_card = QFrame(); io_card.setObjectName("Card")
        io_layout = QVBoxLayout(io_card)
        
        i_layout = QHBoxLayout()
        self.btn_input = QPushButton("")
        self.btn_input.setMinimumHeight(40)
        self.btn_input.setMinimumWidth(160)
        self.btn_input.clicked.connect(self.browse_input)
        self.entry_input = QLineEdit()
        self.entry_input.setMinimumHeight(40)
        i_layout.addWidget(self.btn_input)
        i_layout.addWidget(self.entry_input)
        
        o_layout = QHBoxLayout()
        self.btn_output = QPushButton("")
        self.btn_output.setMinimumHeight(40)
        self.btn_output.setMinimumWidth(160)
        self.btn_output.clicked.connect(self.browse_output)
        self.entry_output = QLineEdit()
        self.entry_output.setMinimumHeight(40)
        o_layout.addWidget(self.btn_output)
        o_layout.addWidget(self.entry_output)
        
        io_layout.addLayout(i_layout)
        io_layout.addSpacing(10)
        io_layout.addLayout(o_layout)
        content_layout.addWidget(io_card)
        
        # OPTIONS
        # Merge Option (Primary/New line)
        self.chk_merge_pages = QCheckBox("")
        self.chk_merge_pages.setFont(QFont(".AppleSystemUIFont", 12, QFont.Weight.Bold))
        self.chk_merge_pages.toggled.connect(self.update_options)
        
        self.lbl_merge_desc = QLabel("")
        self.lbl_merge_desc.setWordWrap(True)
        self.lbl_merge_desc.setStyleSheet("color: #666; font-size: 12px; margin-left: 28px; margin-top: 2px;")
        
        content_layout.addWidget(self.chk_merge_pages)
        content_layout.addWidget(self.lbl_merge_desc)

        opt_layout = QHBoxLayout()
        self.chk_solve = QCheckBox("")
        self.chk_solve.toggled.connect(self.update_options)
        self.chk_cover = QCheckBox("")
        self.chk_cover.toggled.connect(self.update_options)
        
        opt_layout.addWidget(self.chk_solve)
        opt_layout.addWidget(self.chk_cover)
        opt_layout.addStretch()
        content_layout.addLayout(opt_layout)
        
        # SPEED CARD
        speed_card = QFrame(); speed_card.setObjectName("Card")
        speed_v_layout = QVBoxLayout(speed_card)
        
        speed_h_layout = QHBoxLayout()
        self.lbl_speed_title = QLabel("")
        self.lbl_speed_title.setFont(QFont(".AppleSystemUIFont", 11, QFont.Weight.Bold))
        self.combo_speed = QComboBox()
        self.combo_speed.addItems(["Eco (Safe / Chậm)", "Balanced (Standard / Cân bằng)", "Turbo (Max Speed / Nhanh)"])
        self.combo_speed.setMinimumHeight(35)
        self.combo_speed.currentIndexChanged.connect(self.on_speed_changed)
        
        speed_h_layout.addWidget(self.lbl_speed_title)
        speed_h_layout.addWidget(self.combo_speed, 1)
        
        self.lbl_speed_desc = QLabel("")
        self.lbl_speed_desc.setStyleSheet("color: #666; font-size: 12px;")
        self.lbl_speed_desc.setWordWrap(True)
        
        speed_v_layout.addLayout(speed_h_layout)
        speed_v_layout.addWidget(self.lbl_speed_desc)
        content_layout.addWidget(speed_card)
        
        # BUTTONS
        btn_layout = QHBoxLayout()
        self.btn_start = QPushButton("")
        self.btn_start.setObjectName("Primary")
        self.btn_start.setMinimumHeight(60)
        self.btn_start.setFont(QFont(".AppleSystemUIFont", 16, QFont.Weight.Bold))
        self.btn_start.clicked.connect(self.start_processing)
        
        self.btn_stop = QPushButton("")
        self.btn_stop.setObjectName("Danger")
        self.btn_stop.setMinimumHeight(60)
        self.btn_stop.setFont(QFont(".AppleSystemUIFont", 16, QFont.Weight.Bold))
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self.stop_processing)
        
        btn_layout.addWidget(self.btn_start)
        btn_layout.addWidget(self.btn_stop)
        content_layout.addLayout(btn_layout)

        # TIMER
        self.lbl_timer = QLabel("")
        self.lbl_timer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_timer.setFont(QFont(".AppleSystemUIFont", 14, QFont.Weight.Bold))
        self.lbl_timer.setStyleSheet("color: #3a7ebf;")
        content_layout.addWidget(self.lbl_timer)

        # PROGRESS BAR
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setFormat("%v/%m")
        self.progress_bar.hide()
        content_layout.addWidget(self.progress_bar)

        # LOG
        self.log_box = QTextEdit()
        self.log_box.setReadOnly(True)
        self.log_box.setFont(QFont("Menlo", 13))
        self.log_box.setObjectName("Card")
        content_layout.addWidget(self.log_box)

        # FOOTER
        self.lbl_footer = QLabel("")
        self.lbl_footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_footer.setStyleSheet("color: gray;")
        content_layout.addWidget(self.lbl_footer)
        
        self.load_saved_api_key()

    def update_options(self):
        self.solve_var = self.chk_solve.isChecked()
        self.cover_var = self.chk_cover.isChecked()
        self.merge_pages_var = self.chk_merge_pages.isChecked()
        self.save_config()

    def on_speed_changed(self, index):
        self.speed_idx = index
        self.save_config()
        self.update_speed_description()

    def update_speed_description(self):
        desc_keys = ["desc_eco", "desc_balanced", "desc_turbo"]
        self.lbl_speed_desc.setText(self.t(desc_keys[self.speed_idx]))

    def update_ui_texts(self):
        self.setWindowTitle(self.t("title") + " v" + CURRENT_VERSION)
        self.lbl_toolbar.setText(self.t("toolbar"))
        self.btn_merge.setText(self.t("merge_pdf"))
        self.btn_unfinished.setText(self.t("btn_unfinished"))
        self.btn_update.setText(self.t("btn_update"))
        
        self.btn_mode_single.setText(self.t("mode_single"))
        self.btn_mode_batch.setText(self.t("mode_batch"))
        
        self.lbl_api.setText(self.t("api_key"))
        self.entry_api.setPlaceholderText(self.t("api_placeholder"))
        self.btn_load_api.setText(self.t("load_api"))
        
        self.chk_solve.setText(self.t("solve_opt"))
        self.chk_solve.setChecked(self.solve_var)
        self.chk_cover.setText(self.t("cover_opt"))
        self.chk_cover.setChecked(self.cover_var)
        self.chk_merge_pages.setText(self.t("merge_opt"))
        self.chk_merge_pages.setChecked(self.merge_pages_var)
        self.lbl_merge_desc.setText(self.t("merge_desc"))
        
        self.lbl_speed_title.setText(self.t("speed_label"))
        
        self.combo_speed.blockSignals(True)
        self.combo_speed.setCurrentIndex(self.speed_idx)
        self.combo_speed.blockSignals(False)
        self.update_speed_description()

        self.btn_start.setText(self.t("start"))
        self.btn_stop.setText(self.t("stop"))
        
        self.btn_output.setText(self.t("output_dir"))
        self.entry_output.setPlaceholderText(self.t("output_ph"))
        
        self.change_mode(self.mode_group.checkedButton())
        
        self.lbl_timer.setText(self.t("timer_init"))
        self.lbl_footer.setText(self.t("dev_footer", CURRENT_VERSION, self.build_str, self.year_str))
        
        if self.log_box.toPlainText() == "":
            self.write_log(self.t("log_ready"))

    def change_language(self, index):
        new_lang = "VN" if index == 1 else "EN"
        if new_lang != self.current_lang:
            self.current_lang = new_lang
            self.save_config()
            self.update_ui_texts()

    def change_mode(self, btn):
        if not btn: return
        self.entry_input.clear()
        if btn == self.btn_mode_single:
            self.btn_input.setText(self.t("select_pdf"))
            self.entry_input.setPlaceholderText(self.t("input_pdf_ph"))
        else:
            self.btn_input.setText(self.t("input_dir_btn"))
            self.entry_input.setPlaceholderText(self.t("input_dir_ph"))

    def closeEvent(self, event):
        key = self.entry_api.text().strip()
        if key: self.save_api_key(key)
        self.stop_event.set()
        event.accept()

    def save_api_key(self, key):
        try:
            if not os.path.exists(CONFIG_DIR): os.makedirs(CONFIG_DIR)
            with open(API_KEY_FILE, "w", encoding="utf-8") as f: f.write(key)
        except: pass

    def load_saved_api_key(self):
        if os.path.exists(API_KEY_FILE):
            try:
                with open(API_KEY_FILE, "r", encoding="utf-8") as f:
                    key = f.read().strip()
                    if key: self.entry_api.setText(key)
            except: pass

    def load_api_from_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Load API Key", "", "Text Files (*.txt)")
        if path:
            try:
                with open(path, "r") as f: self.entry_api.setText(f.read().strip())
                self.write_log("[+] API Key loaded.")
            except: pass

    def write_log(self, text):
        self.log_box.append(text)
        scroll = self.log_box.verticalScrollBar()
        scroll.setValue(scroll.maximum())

    def browse_input(self):
        if self.btn_mode_single.isChecked():
            path, _ = QFileDialog.getOpenFileName(self, self.t("select_pdf"), "", "PDF (*.pdf)")
        else:
            path = QFileDialog.getExistingDirectory(self, self.t("input_dir_btn"))
        if path: self.entry_input.setText(path)

    def browse_output(self):
        path = QFileDialog.getExistingDirectory(self, self.t("output_dir"))
        if path: self.entry_output.setText(path)

    def stop_processing(self):
        self.write_log(self.t("log_stop_cmd"))
        self.stop_event.set()
        self.btn_stop.setEnabled(False)

    def start_processing(self):
        api_key = self.entry_api.text().strip()
        input_path = self.entry_input.text().strip()
        output_dir = self.entry_output.text().strip()
        
        if not api_key:
            self.write_log(self.t("err_api"))
            return
        if not input_path or not os.path.exists(input_path):
            self.write_log(self.t("err_input"))
            return

        if not output_dir:
            if os.path.isdir(input_path):
                output_dir = input_path
            else:
                output_dir = os.path.dirname(input_path) 
            self.write_log(self.t("msg_auto_out", output_dir))
        elif not os.path.exists(output_dir):
            self.write_log(self.t("err_output"))
            return

        self.save_api_key(api_key)
        self.start_time = datetime.now()
        self.timer_running = True
        self.timer.start(1000)

        self.btn_start.setEnabled(False)
        self.btn_input.setEnabled(False)
        self.btn_output.setEnabled(False)
        self.btn_mode_single.setEnabled(False)
        self.btn_mode_batch.setEnabled(False)
        self.btn_stop.setEnabled(True)
        self.stop_event.clear()
        
        self.client = genai.Client(api_key=api_key)
        
        mode = self.t("mode_single") if self.btn_mode_single.isChecked() else self.t("mode_batch")
        
        resume_at = 0
        initial_content = ""

        # Check checkpoint for Single Mode solely to prompt
        if mode == self.t("mode_single"):
            abs_input = os.path.abspath(input_path)
            pdf_id = hashlib.md5(abs_input.encode('utf-8')).hexdigest()
            cp_path = os.path.join(CHECKPOINT_DIR, f"{pdf_id}.json")
            if os.path.exists(cp_path):
                try:
                    with open(cp_path, "r", encoding="utf-8") as f:
                        cp_data = json.load(f)
                        page = cp_data.get("current_page", 0)
                        if page > 0:
                            ret = QMessageBox.question(self, self.t("resume_title"), 
                                                     self.t("resume_query", page),
                                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                            if ret == QMessageBox.StandardButton.Yes:
                                resume_at = page
                                initial_content = cp_data.get("content", "")
                            else:
                                # User chose No, clear checkpoint to start fresh
                                if os.path.exists(cp_path): os.remove(cp_path)
                except: pass

        self.worker = WorkerThread(self, input_path, output_dir, mode, resume_at, initial_content, self.speed_idx, self.client)
        self.worker.log_signal.connect(self.write_log)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.status_signal.connect(self.handle_status)
        self.worker.finished_signal.connect(self.reset_ui)
        
        self.progress_bar.show()
        self.progress_bar.setValue(0)
        self.progress_bar.setObjectName("")
        self.progress_bar.setStyleSheet("")
        
        self.worker.start()

    def handle_status(self, msg, is_error):
        if is_error:
            self.progress_bar.setObjectName("Error")
            self.progress_bar.setFormat(msg)
            self.progress_bar.setStyleSheet("QProgressBar#Error { color: white; } QProgressBar#Error::chunk { background-color: #a83232; }")
        else:
            self.progress_bar.setObjectName("")
            self.progress_bar.setFormat("%v/%m")
            self.progress_bar.setStyleSheet("")

    def update_progress(self, current, total):
        if total > 0:
            self.progress_bar.setMaximum(total)
            self.progress_bar.setValue(current)

    def reset_ui(self):
        self.timer_running = False
        self.timer.stop()
        self.btn_start.setEnabled(True)
        self.btn_input.setEnabled(True)
        self.btn_output.setEnabled(True)
        self.btn_mode_single.setEnabled(True)
        self.btn_mode_batch.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.progress_bar.hide()

    def update_timer_ui(self):
        if self.timer_running and self.start_time:
            delta = datetime.now() - self.start_time
            m, s = divmod(int(delta.total_seconds()), 60)
            self.lbl_timer.setText(self.t("timer", m, s))

    def check_for_updates(self):
        try:
            req = urllib.request.Request(GITHUB_API_URL, headers={'User-Agent': 'PDFScan2Word-App'})
            with urllib.request.urlopen(req, timeout=5) as r:
                data = json.loads(r.read().decode('utf-8'))
                latest = data.get('tag_name', '').lstrip('v')
            if latest > CURRENT_VERSION:
                if sys.platform == "win32":
                    suffix = "Windows.zip"
                elif sys.platform == "darwin":
                    arch = platform.machine()
                    if arch == "arm64":
                        suffix = "macOS-arm64.zip"
                    else:
                        suffix = "macOS-x86_64.zip"
                else:
                    self.write_log(f"\\n[*] Cập nhật mới có sẵn: v{latest} - Vui lòng tải tại Github.")
                    return
                    
                download_url = ""
                for asset in data.get('assets', []):
                    if asset.get('name', '').endswith(suffix):
                        download_url = asset.get('browser_download_url', '')
                        break
                        
                self.latest_version = latest
                self.download_url = download_url
                
                if download_url:
                    self.update_available_signal.emit(latest, download_url)
                else:
                    self.write_log(f"\\n[*] Cập nhật mới có sẵn: v{latest} - Vui lòng tải tại Github.")
        except: pass

    def prompt_update(self, latest, download_url):
        title = self.t("update_title")
        msg = self.t("update_msg", latest, CURRENT_VERSION)
        
        box = QMessageBox(self)
        box.setWindowTitle(title)
        box.setText(msg)
        box.setIcon(QMessageBox.Icon.Information)
        btn_yes = box.addButton(self.t("btn_yes"), QMessageBox.ButtonRole.AcceptRole)
        btn_no = box.addButton(self.t("btn_no"), QMessageBox.ButtonRole.RejectRole)
        box.exec()
        
        if box.clickedButton() == btn_yes:
            dlg = UpdateProgressDialog(self, download_url, latest)
            dlg.exec()

    def perform_update_swap(self, temp_dir):
        if not getattr(sys, 'frozen', False):
            QMessageBox.warning(self, "Lỗi", "Tính năng tự động cập nhật chỉ hoạt động trong bản release (.exe/.app).\\nVui lòng cập nhật mã nguồn bằng Git.")
            return

        current_exe = sys.executable
        if sys.platform == "darwin":
            app_path = os.path.dirname(os.path.dirname(os.path.dirname(current_exe)))
            if app_path.endswith('.app'):
                target_path = app_path
            else:
                target_path = current_exe
        else:
            target_path = current_exe
            
        new_app_path = None
        for item in os.listdir(temp_dir):
            if item.endswith(".app") or item.endswith(".exe"):
                new_app_path = os.path.join(temp_dir, item)
                break
                
        if not new_app_path:
            QMessageBox.critical(self, "Lỗi cập nhật", "Không tìm thấy file chạy (.exe/.app) trong file tải về!")
            return
            
        if sys.platform == "win32":
            script_path = os.path.join(temp_dir, "update.bat")
            with open(script_path, "w", encoding="utf-8") as f:
                f.write(f'''@echo off
timeout /t 2 /nobreak > NUL
move /y "{new_app_path}" "{target_path}"
start "" "{target_path}"
del "%~f0"
''')
            subprocess.Popen([script_path], creationflags=subprocess.CREATE_NO_WINDOW)
        else:
            script_path = os.path.join(temp_dir, "update.sh")
            with open(script_path, "w", encoding="utf-8") as f:
                f.write(f'''#!/bin/bash
sleep 2
rm -rf "{target_path}"
mv "{new_app_path}" "{target_path}"
xattr -cr "{target_path}" 2>/dev/null
open "{target_path}"
rm "$0"
''')
            os.chmod(script_path, stat.S_IRWXU)
            subprocess.Popen([script_path])
            
        QApplication.quit()

    def open_unfinished_manager(self):
        dlg = CheckpointHistoryDialog(self)
        if dlg.exec():
            data = dlg.selected_data
            opts = data.get("options", {})
            
            # 1. Update Mode (Quan trọng: Phải làm trước khi update_ui_texts)
            saved_mode = opts.get("mode")
            if saved_mode == self.t("mode_batch"):
                self.btn_mode_batch.setChecked(True)
            else:
                self.btn_mode_single.setChecked(True)
            
            # 2. Update Options
            self.chk_solve.setChecked(opts.get("solve", False))
            self.chk_cover.setChecked(opts.get("cover", False))
            self.chk_merge_pages.setChecked(opts.get("merge", False))
            self.current_lang = opts.get("lang", "VN")
            
            # 3. Update Language Menu
            idx = 0 if self.current_lang == "EN" else 1
            self.lang_menu.setCurrentIndex(idx)
            
            # 4. Update UI Texts (Cái này sẽ gọi change_mode và clear entry_input)
            self.update_ui_texts()
            
            # 5. Set Path (Cuối cùng để không bị clear)
            pdf_path = data.get("pdf_path", "")
            self.entry_input.setText(pdf_path)
            
            self.write_log(f"[*] Loaded checkpoint for: {os.path.basename(pdf_path)}")

    def auto_check_checkpoints(self):
        if not os.path.exists(CHECKPOINT_DIR): return
        files = [f for f in os.listdir(CHECKPOINT_DIR) if f.endswith(".json")]
        if files:
            ret = QMessageBox.information(self, self.t("resume_title"), 
                                        f"Phát hiện {len(files)} tiến trình chưa hoàn thành. Bạn có muốn xem danh sách để tiếp tục không?\nDetected {len(files)} unfinished tasks. View them to resume?",
                                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if ret == QMessageBox.StandardButton.Yes:
                self.open_unfinished_manager()

    def open_merge_popup(self):
        self.merge_window = MergeWindow(self)
        self.merge_window.exec()

    def open_update_menu(self):
        dlg = UpdateMenuDialog(self)
        dlg.exec()

def _qt_message_handler(mode, context, message):
    msg = str(message)
    if "accessibilityLabel on invalid object" in msg: return
    if "Populating font family aliases" in msg: return
    if "Could not parse stylesheet" in msg: return
    print(msg)

if __name__ == "__main__":
    from PyQt6.QtCore import qInstallMessageHandler
    qInstallMessageHandler(_qt_message_handler)
    
    app = QApplication(sys.argv)
    window = PDFOCRApp()
    window.show()
    sys.exit(app.exec())
